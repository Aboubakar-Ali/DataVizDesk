from docx import Document
from docx.shared import Inches
from tkinter import filedialog, simpledialog, messagebox
import re
import os

def extract_metadata(filepath):
    with open(filepath, 'r', encoding='utf-8') as file:
        text = file.read()
    
    title_search = re.search(r'Title:\s*(.+)', text)
    author_search = re.search(r'Author:\s*(.+)', text)
    
    title = title_search.group(1).strip() if title_search else 'Titre inconnu'
    author = author_search.group(1).strip() if author_search else 'Auteur inconnu'
    
    return title, author

def create_title_page(doc_path, title, author, image_path, report_author):
    doc = Document()
    doc.add_heading(title, 0)
    doc.add_picture(image_path, width=Inches(4.0))
    doc.add_paragraph(f'Auteur du Livre: {author}')
    doc.add_paragraph(f'Auteur du Rapport: {report_author}')
    doc.save(doc_path)

def generate_report(window):
    report_author = simpledialog.askstring("Nom", "Entrez votre nom pour le rapport:", parent=window)
    if not report_author:
        messagebox.showwarning("Annulation", "Génération du rapport annulée.")
        return

    book_filepath = filedialog.askopenfilename(title="Sélectionnez le fichier du livre", parent=window)
    image_path = filedialog.askopenfilename(title="Sélectionnez l'image pour le rapport", parent=window)

    if not book_filepath or not image_path:
        messagebox.showwarning("Annulation", "Génération du rapport annulée.")
        return

    title, author = extract_metadata(book_filepath)
    
    save_directory = 'Data Explorration/Doc'
    if not os.path.exists(save_directory):
        os.makedirs(save_directory)
    doc_path = os.path.join(save_directory, 'rapport.docx')
    
    create_title_page(doc_path, title, author, image_path, report_author)
    messagebox.showinfo("Succès", f"Le document Word a été créé avec succès et enregistré à l'emplacement : {doc_path}")
